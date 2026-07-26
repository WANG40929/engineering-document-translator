from __future__ import annotations

import time

from docx import Document

from ..i18n import tr
from ..models import FileResult
from ..text_utils import is_translatable
from .base import TranslationEngine


def _table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _table_paragraphs(nested)


def _all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        yield from _table_paragraphs(table)
    for section in document.sections:
        for part in (section.header, section.footer, section.first_page_header, section.first_page_footer, section.even_page_header, section.even_page_footer):
            yield from part.paragraphs
            for table in part.tables:
                yield from _table_paragraphs(table)


class DocxEngine(TranslationEngine):
    extensions = (".docx",)

    def translate(self, source, destination, translator, options, progress=None) -> FileResult:
        started = time.monotonic()
        result = FileResult(str(source), str(destination), engine="DOCX")
        try:
            document = Document(source)
            paragraphs = list(_all_paragraphs(document))
            units = []
            for paragraph in paragraphs:
                text_runs = [run for run in paragraph.runs if run.text]
                text = "".join(run.text for run in text_runs)
                if text_runs and is_translatable(text):
                    units.append((text_runs, text))
            texts = [text for _runs, text in units]

            def report(done, total):
                if progress:
                    progress(
                        str(source),
                        done / max(total, 1),
                        tr("progress.word", done=done, total=total),
                    )

            translated = translator.translate_many(texts, report)
            for (runs, _source_text), value in zip(units, translated):
                # Keep the paragraph/table cell itself and its leading character style.
                # Sending the complete paragraph gives the model enough context for
                # split runs and bilingual labels; secondary runs are emptied only of text.
                runs[0].text = value
                for run in runs[1:]:
                    run.text = ""
            result.translated_units = len(units)
            result.skipped_units = sum(1 for p in paragraphs if p.text and not is_translatable(p.text))
            destination.parent.mkdir(parents=True, exist_ok=True)
            document.save(destination)
            result.status = "completed"
        except Exception as exc:
            result.status = "failed"
            result.errors.append(str(exc))
            if destination.exists():
                destination.unlink()
        result.elapsed_seconds = round(time.monotonic() - started, 2)
        result.usage = dict(getattr(translator, "usage", {}))
        return result

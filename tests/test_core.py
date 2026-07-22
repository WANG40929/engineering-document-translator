from __future__ import annotations

import csv
import json
import sqlite3
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import fitz
from docx import Document
from openpyxl import Workbook, load_workbook

from translator_app.cache import TranslationCache
from translator_app.config import ConfigStore
from translator_app.deepseek import DeepSeekError, DeepSeekTranslator, IdentityTranslator, IncompleteResponseError
from translator_app.engines.csv_engine import CsvEngine
from translator_app.engines.docx_engine import DocxEngine
from translator_app.engines.babeldoc_engine import BabelDocEngine
from translator_app.engines.pdf_engine import PdfEngine
from translator_app.engines.xlsx_engine import XlsxEngine
from translator_app.models import TranslationOptions
from translator_app.pipeline import TranslationPipeline
from translator_app.secret_store import SecretStore
from translator_app.text_utils import is_translatable, protect_text


class PrefixTranslator:
    usage = {"offline": True}

    def translate_many(self, texts, progress=None):
        if progress:
            progress(len(texts), len(texts))
        return [f"中译：{text}" for text in texts]


class CoreTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory(prefix="udt_test_")
        self.root = Path(self.temp.name)
        self.options = TranslationOptions(target_language="zh")

    def tearDown(self):
        self.temp.cleanup()

    def test_token_protection_and_classification(self):
        protected = protect_text("Pump KZ5001-MB-010 is 6.3 kV")
        self.assertIn("__UDT_", protected.text)
        self.assertEqual(protected.restore(protected.text), "Pump KZ5001-MB-010 is 6.3 kV")
        self.assertEqual(protected.restore("Pump UDT_0000 is UDT 0001"), "Pump KZ5001-MB-010 is 6.3 kV")
        self.assertFalse(is_translatable("KZ5001-MB-010"))
        self.assertTrue(is_translatable("Lube oil pump"))

    def test_api_batch_deduplicates_and_caches(self):
        class MockDeepSeek(DeepSeekTranslator):
            calls = 0
            def _request(inner, items):
                inner.calls += 1
                return {item["id"]: "译文" for item in items}

        cache = TranslationCache(self.root / "cache.sqlite3")
        translator = MockDeepSeek("test-key", cache=cache)
        self.assertEqual(translator.translate_many(["Bearing", "Bearing", "Pump"]), ["译文", "译文", "译文"])
        self.assertEqual(translator.calls, 1)
        translator.translate_many(["Bearing", "Pump"])
        self.assertEqual(translator.calls, 1)

    def test_api_repairs_only_missing_segments(self):
        class PartialDeepSeek(DeepSeekTranslator):
            calls = []

            def _request(inner, items, review=False):
                ids = [int(item["id"]) for item in items]
                inner.calls.append(ids)
                if len(inner.calls) == 1:
                    missing = set(range(11, 22))
                    partial = {item_id: f"译文{item_id}" for item_id in ids if item_id not in missing}
                    raise IncompleteResponseError(missing, partial, "stop")
                return {item_id: f"译文{item_id}" for item_id in ids}

        translator = PartialDeepSeek("test-key", cache=TranslationCache(self.root / "repair.sqlite3"))
        result = translator.translate_many([f"Segment {index} text" for index in range(29)])
        self.assertEqual(len(result), 29)
        self.assertEqual(translator.calls[1], list(range(11, 22)))
        self.assertEqual(translator.usage["repair_requests"], 1)
        self.assertEqual(translator.usage["recovered_segments"], 18)

    def test_api_splits_an_invalid_batch_down_to_single_segments(self):
        class SplitDeepSeek(DeepSeekTranslator):
            def _request(inner, items, review=False):
                ids = {int(item["id"]) for item in items}
                if len(items) > 1:
                    raise IncompleteResponseError(ids, {}, "stop")
                return {int(items[0]["id"]): "单段译文"}

        translator = SplitDeepSeek("test-key", cache=TranslationCache(self.root / "split.sqlite3"))
        self.assertEqual(translator.translate_many(["First segment", "Second segment"]), ["单段译文", "单段译文"])
        self.assertGreaterEqual(translator.usage["split_retries"], 1)

    def test_api_disables_thinking_and_records_finish_reason(self):
        class FakeResponse:
            def __enter__(self): return self
            def __exit__(self, *_args): return False
            def read(self):
                return json.dumps({
                    "choices": [{
                        "finish_reason": "stop",
                        "message": {"content": '{"translations":[{"id":0,"text":"译文"}]}'},
                    }],
                    "usage": {"prompt_tokens": 5, "completion_tokens": 3, "total_tokens": 8},
                }).encode("utf-8")

        translator = DeepSeekTranslator("test-key", cache=TranslationCache(self.root / "payload.sqlite3"))
        with patch("translator_app.deepseek.urllib.request.urlopen", return_value=FakeResponse()) as mocked:
            self.assertEqual(translator._request([{"id": 0, "text": "Pump"}]), {0: "译文"})
        request = mocked.call_args.args[0]
        payload = json.loads(request.data.decode("utf-8"))
        self.assertEqual(payload["thinking"], {"type": "disabled"})
        self.assertGreaterEqual(payload["max_tokens"], 2048)
        self.assertEqual(translator.usage["finish_reasons"], {"stop": 1})

    def test_api_retries_a_hallucinated_placeholder(self):
        class FakeResponse:
            def __init__(self, text): self.text = text
            def __enter__(self): return self
            def __exit__(self, *_args): return False
            def read(self):
                content = json.dumps({"translations": [{"id": 0, "text": self.text}]}, ensure_ascii=False)
                return json.dumps({
                    "choices": [{"finish_reason": "stop", "message": {"content": content}}],
                    "usage": {},
                }, ensure_ascii=False).encode("utf-8")

        translator = DeepSeekTranslator(
            "test-key",
            cache=TranslationCache(self.root / "placeholder-response.sqlite3"),
            quality_review=False,
        )
        responses = [FakeResponse("安全 UDT_0000"), FakeResponse("安全要求")]
        with patch("translator_app.deepseek.urllib.request.urlopen", side_effect=responses) as mocked:
            self.assertEqual(translator.translate_many(["Safety requirements"]), ["安全要求"])
        self.assertEqual(mocked.call_count, 2)
        self.assertEqual(translator.usage["schema_failures"], 1)

    def test_polluted_cache_is_repaired_or_selectively_refreshed(self):
        class RepairingDeepSeek(DeepSeekTranslator):
            calls = []
            def _request(inner, items, review=False):
                inner.calls.extend(item["text"] for item in items)
                return {int(item["id"]): "安全要求" for item in items}

        cache = TranslationCache(self.root / "polluted-cache.sqlite3")
        translator = RepairingDeepSeek("test-key", cache=cache, quality_review=False)
        code_source = "Pump SEPCO1 is ready"
        plain_source = "Safety requirements"
        cache.put_many("auto", "zh", [
            (code_source, "泵 UDT_0000 已就绪"),
            (plain_source, "安全 UDT_0000"),
        ], translator._cache_signature)

        result = translator.translate_many([code_source, plain_source])
        self.assertEqual(result, ["泵 SEPCO1 已就绪", "安全要求"])
        self.assertEqual(len(translator.calls), 1)
        self.assertIn("Safety requirements", translator.calls[0])
        self.assertNotIn(
            "UDT_",
            cache.get("auto", "zh", plain_source, translator._cache_signature),
        )

    def test_cache_trims_oldest_rows_and_compacts_file(self):
        path = self.root / "limited-cache.sqlite3"
        cache = TranslationCache(path, max_cache_bytes=64 * 1024)
        pairs = [(f"source-{index}-" + "x" * 600, f"target-{index}-" + "y" * 600) for index in range(200)]
        cache.put_many("en", "zh", pairs)
        conn = sqlite3.connect(path)
        try:
            remaining = conn.execute("SELECT COUNT(*) FROM translations").fetchone()[0]
            oldest = conn.execute("SELECT source_text FROM translations ORDER BY rowid LIMIT 1").fetchone()[0]
        finally:
            conn.close()
        self.assertLess(remaining, len(pairs))
        self.assertNotIn("source-0-", oldest)
        self.assertLessEqual(path.stat().st_size, 64 * 1024)

    def test_legacy_model_config_is_migrated(self):
        path = self.root / "config.json"
        path.write_text('{"model":"deepseek-chat","target_language":"zh"}', encoding="utf-8")
        config = ConfigStore(path).load()
        self.assertEqual(config.model, "deepseek-v4-flash")
        self.assertFalse(config.pure_target_language)

    def test_pure_target_choice_is_respected_after_migration(self):
        path = self.root / "config-v2.json"
        path.write_text('{"config_version":2,"pure_target_language":true}', encoding="utf-8")
        self.assertTrue(ConfigStore(path).load().pure_target_language)

    def test_v2_config_gains_safe_pdf_defaults(self):
        path = self.root / "config-v2-pdf.json"
        path.write_text('{"config_version":2,"model":"deepseek-v4-flash"}', encoding="utf-8")
        config = ConfigStore(path).load()
        self.assertEqual(config.pdf_mode, "auto")
        self.assertEqual(config.pdf_output, "mono")
        self.assertEqual(config.config_version, 3)

    def test_smart_pdf_backend_generates_mono_and_dual_without_key_in_argv(self):
        source, output = self.root / "report.pdf", self.root / "report_ZH.pdf"
        document = fitz.open()
        document.new_page(width=400, height=300).insert_text((30, 50), "Long report paragraph for translation", fontsize=12)
        document.save(source)
        document.close()
        fake = self.root / "fake_babeldoc.py"
        fake.write_text(
            """import argparse
from pathlib import Path
import fitz
p = argparse.ArgumentParser()
p.add_argument('--config'); p.add_argument('--files'); p.add_argument('--output')
p.add_argument('--working-dir'); p.add_argument('--lang-in'); p.add_argument('--lang-out')
a, rest = p.parse_known_args()
assert 'sk-super-secret' not in ' '.join(__import__('sys').argv)
assert 'sk-super-secret' in Path(a.config).read_text(encoding='utf-8')
out = Path(a.output); out.mkdir(parents=True, exist_ok=True)
for kind in ('mono', 'dual'):
    d = fitz.open(); d.new_page(width=300, height=200).insert_text((20, 40), kind)
    d.save(out / f'sample.no_watermark.zh.{kind}.pdf'); d.close()
print('Translate Paragraphs 65%')
print('Save PDF 100%')
""",
            encoding="utf-8",
        )

        class SmartTranslator:
            api_key = "sk-super-secret"
            base_url = "https://api.deepseek.com/chat/completions"
            usage = {}

        options = TranslationOptions(
            target_language="zh", model="deepseek-v4-flash", pdf_mode="smart",
            pdf_output="both", babeldoc_path=fake,
        )
        updates = []
        result = BabelDocEngine().translate(
            source, output, SmartTranslator(), options,
            lambda _file, fraction, message: updates.append((fraction, message)),
        )
        self.assertEqual(result.status, "completed", result.errors)
        self.assertTrue(output.exists())
        self.assertEqual(len(result.additional_outputs), 1)
        self.assertTrue(Path(result.additional_outputs[0]).exists())
        self.assertEqual(updates[-1][0], 1.0)

    def test_auto_pdf_mode_uses_smart_only_for_prose_when_available(self):
        pipeline = TranslationPipeline()
        source = self.root / "sample.pdf"
        source.touch()
        options = TranslationOptions(pdf_mode="auto", babeldoc_path=Path(sys.executable))
        with patch.object(BabelDocEngine, "looks_like_prose", return_value=True):
            self.assertIs(pipeline.engine_for(source, options), pipeline.smart_pdf_engine)
        with patch.object(BabelDocEngine, "looks_like_prose", return_value=False):
            self.assertIs(pipeline.engine_for(source, options), pipeline.strict_pdf_engine)

    def test_babeldoc_progress_maps_stages_and_rich_ratio(self):
        self.assertGreaterEqual(
            BabelDocEngine._progress_from_output("Translate Paragraphs", 0.1), 0.48
        )
        self.assertEqual(
            BabelDocEngine._progress_from_output("translate 73.0/100", 0.2), 0.73
        )

    def test_babeldoc_glossary_csv_has_required_columns(self):
        path = self.root / "glossary.csv"
        BabelDocEngine._write_glossary(path, {"bearing": "轴承"}, "zh")
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.DictReader(handle))
        self.assertEqual(rows, [{"source": "bearing", "target": "轴承", "tgt_lng": "zh"}])

    def test_quality_review_retries_real_words_but_not_short_codes(self):
        class ReviewDeepSeek(DeepSeekTranslator):
            reviews = []
            def _request(inner, items, review=False):
                if review:
                    inner.reviews.extend(item["text"] for item in items)
                    return {item["id"]: "六角螺母" for item in items}
                return {item["id"]: item["text"] for item in items}

        translator = ReviewDeepSeek("test-key", cache=TranslationCache(self.root / "review.sqlite3"))
        result = translator.translate_many(["HEXAGON NUT", "PCE"])
        self.assertEqual(result, ["六角螺母", "PCE"])
        self.assertEqual(translator.reviews, ["HEXAGON NUT"])
        self.assertEqual(translator.usage["quality_retries"], 1)

    def test_dpapi_secret_round_trip(self):
        store = SecretStore(self.root / "key.bin")
        store.save("sk-test-123")
        self.assertEqual(store.load(), "sk-test-123")

    def test_pdf_text_layer_and_blank_page(self):
        source, output = self.root / "sample.pdf", self.root / "sample_ZH.pdf"
        document = fitz.open()
        page = document.new_page(width=400, height=300)
        page.draw_line((20, 100), (380, 100), color=(1, 0, 0), width=2)
        page.insert_text((30, 50), "Lube oil pump KZ5001-MB-010", fontsize=12)
        document.new_page(width=400, height=300)
        document.save(source); document.close()
        result = PdfEngine().translate(source, output, IdentityTranslator(), self.options)
        self.assertEqual(result.status, "completed")
        self.assertEqual(result.skipped_pages, [2])
        translated = fitz.open(output)
        self.assertEqual(translated.page_count, 2)
        self.assertIn("Lube oil pump", translated[0].get_text())
        self.assertTrue(translated[0].get_drawings())
        translated.close()

    def test_pdf_large_centered_title_is_not_lost(self):
        source, output = self.root / "title.pdf", self.root / "title_ZH.pdf"
        document = fitz.open()
        page = document.new_page(width=400, height=300)
        page.insert_text(
            (45, 70),
            "Lube Oil Pump Installation Manual",
            fontname="helv",
            fontsize=22,
        )
        document.save(source); document.close()

        class TitleTranslator:
            usage = {"offline": True}
            def translate_many(self, texts, progress=None):
                return ["润滑油泵安装调试手册" for _ in texts]

        result = PdfEngine().translate(source, output, TitleTranslator(), self.options)
        self.assertEqual(result.status, "completed")
        translated = fitz.open(output)
        try:
            self.assertIn("润滑油泵安装调试手册", translated[0].get_text())
        finally:
            translated.close()

    def test_pdf_long_body_line_keeps_original_start_position(self):
        source, output = self.root / "body-line.pdf", self.root / "body-line_ZH.pdf"
        document = fitz.open()
        page = document.new_page(width=595.3, height=400)
        page.insert_text(
            (92.8, 180),
            "Still equipment stored inside the building is not adequately protected and may be damaged",
            fontname="helv",
            fontsize=10,
        )
        source_line = page.get_text("dict")["blocks"][0]["lines"][0]
        source_x0 = fitz.Rect(source_line["bbox"]).x0
        document.save(source); document.close()

        class ShortTranslator:
            usage = {"offline": True}
            def translate_many(self, texts, progress=None):
                return ["Equipment is unprotected" for _ in texts]

        result = PdfEngine().translate(source, output, ShortTranslator(), self.options)
        self.assertEqual(result.status, "completed")
        translated = fitz.open(output)
        try:
            translated_line = translated[0].get_text("dict")["blocks"][0]["lines"][0]
            translated_x0 = fitz.Rect(translated_line["bbox"]).x0
            self.assertAlmostEqual(translated_x0, source_x0, delta=1.0)
        finally:
            translated.close()

    def test_pdf_translation_uses_adjacent_whitespace_before_shrinking_font(self):
        source, output = self.root / "header-fields.pdf", self.root / "header-fields_ZH.pdf"
        document = fitz.open()
        page = document.new_page(width=595.3, height=300)
        page.insert_text((414.6, 50), "Page:", fontsize=10)
        page.insert_text((467.9, 50), "2 of 4", fontsize=10)
        document.save(source); document.close()

        class HeaderTranslator:
            usage = {"offline": True}
            def translate_many(self, texts, progress=None):
                return ["Page label", "Page 2 of 4"]

        result = PdfEngine().translate(source, output, HeaderTranslator(), self.options)
        self.assertEqual(result.status, "completed")
        translated = fitz.open(output)
        try:
            spans = [
                span
                for block in translated[0].get_text("dict")["blocks"]
                for line in block.get("lines", [])
                for span in line.get("spans", [])
                if span.get("text", "").strip()
            ]
            self.assertGreaterEqual(min(span["size"] for span in spans), 9.5)
        finally:
            translated.close()

    def test_pdf_standalone_translation_wraps_without_tiny_font(self):
        source, output = self.root / "standalone.pdf", self.root / "standalone_ZH.pdf"
        document = fitz.open()
        page = document.new_page(width=595.3, height=300)
        page.insert_text((56.8, 80), "requirements", fontsize=11)
        document.save(source); document.close()

        class LongTranslator:
            usage = {"offline": True}
            def translate_many(self, texts, progress=None):
                return [
                    "All personnel must comply with the site safety requirements and use approved equipment."
                ]

        result = PdfEngine().translate(source, output, LongTranslator(), self.options)
        self.assertEqual(result.status, "completed")
        translated = fitz.open(output)
        try:
            spans = [
                span
                for block in translated[0].get_text("dict")["blocks"]
                for line in block.get("lines", [])
                for span in line.get("spans", [])
                if span.get("text", "").strip()
            ]
            self.assertTrue(spans)
            self.assertGreaterEqual(min(span["size"] for span in spans), 10.5)
        finally:
            translated.close()

    def test_pdf_failure_reports_exact_page_and_preserves_real_progress(self):
        source = self.root / "two-pages.pdf"
        document = fitz.open()
        document.new_page(width=400, height=300).insert_text((30, 50), "First page text", fontsize=12)
        document.new_page(width=400, height=300).insert_text((30, 50), "Second page text", fontsize=12)
        document.save(source); document.close()

        class FailingTranslator:
            usage = {"offline": True}
            calls = 0
            def translate_many(inner, texts, progress=None):
                inner.calls += 1
                if inner.calls == 2:
                    raise DeepSeekError("simulated missing segment")
                if progress:
                    progress(len(texts), len(texts))
                return list(texts)

        updates = []
        result = PdfEngine().translate(
            source,
            self.root / "failed_ZH.pdf",
            FailingTranslator(),
            self.options,
            lambda _file, fraction, message: updates.append((fraction, message)),
        )
        self.assertEqual(result.status, "failed")
        self.assertEqual(result.translated_units, 1)
        self.assertIn("第 2/2 页", result.errors[0])
        self.assertLess(updates[-1][0], 1.0)

    def test_pipeline_does_not_report_full_progress_when_pdf_fails(self):
        source = self.root / "pipeline-failure.pdf"
        document = fitz.open()
        document.new_page(width=400, height=300).insert_text((30, 50), "First page text", fontsize=12)
        document.new_page(width=400, height=300).insert_text((30, 50), "Second page text", fontsize=12)
        document.save(source); document.close()

        class FailingTranslator:
            usage = {}
            calls = 0
            def translate_many(inner, texts, progress=None):
                inner.calls += 1
                if inner.calls == 2:
                    raise DeepSeekError("simulated failure")
                return list(texts)

        updates = []
        options = TranslationOptions(target_language="zh", output_dir=self.root)
        results = TranslationPipeline().run(
            [source], FailingTranslator(), options,
            lambda _file, fraction, message: updates.append((fraction, message)),
        )
        self.assertEqual(results[0].status, "failed")
        self.assertLess(updates[-1][0], 1.0)
        self.assertIn("失败", updates[-1][1])

    def test_docx_preserves_table(self):
        source, output = self.root / "sample.docx", self.root / "sample_ZH.docx"
        document = Document(); document.add_heading("Installation manual", 1)
        table = document.add_table(rows=1, cols=2); table.cell(0, 0).text = "Equipment"; table.cell(0, 1).text = "Bearing"
        document.save(source)
        result = DocxEngine().translate(source, output, PrefixTranslator(), self.options)
        self.assertEqual(result.status, "completed")
        translated = Document(output)
        self.assertEqual(len(translated.tables), 1)
        self.assertTrue(translated.paragraphs[0].text.startswith("中译："))

    def test_docx_translates_split_runs_as_one_paragraph(self):
        source, output = self.root / "split.docx", self.root / "split_ZH.docx"
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Packliste / ").bold = True
        paragraph.add_run("Packing List")
        document.save(source)

        class ContextTranslator:
            usage = {"offline": True}
            seen = []
            def translate_many(inner, texts, progress=None):
                inner.seen.extend(texts)
                return ["装箱单"]

        translator = ContextTranslator()
        result = DocxEngine().translate(source, output, translator, self.options)
        self.assertEqual(result.status, "completed")
        self.assertEqual(translator.seen, ["Packliste / Packing List"])
        self.assertEqual(Document(output).paragraphs[0].text, "装箱单")

    def test_xlsx_preserves_formula(self):
        source, output = self.root / "sample.xlsx", self.root / "sample_ZH.xlsx"
        book = Workbook(); sheet = book.active; sheet["A1"] = "Cable description"; sheet["B1"] = 5; sheet["C1"] = "=B1*2"; book.save(source)
        result = XlsxEngine().translate(source, output, PrefixTranslator(), self.options)
        self.assertEqual(result.status, "completed")
        translated = load_workbook(output, data_only=False)
        self.assertTrue(translated.active["A1"].value.startswith("中译："))
        self.assertEqual(translated.active["C1"].value, "=B1*2")
        translated.close()

    def test_csv_preserves_shape(self):
        source, output = self.root / "sample.csv", self.root / "sample_ZH.csv"
        with source.open("w", encoding="utf-8", newline="") as handle:
            csv.writer(handle, delimiter=";").writerows([["Equipment", "Value"], ["Pump", "6.3 kV"]])
        result = CsvEngine().translate(source, output, PrefixTranslator(), self.options)
        self.assertEqual(result.status, "completed")
        with output.open("r", encoding="utf-8", newline="") as handle:
            rows = list(csv.reader(handle, delimiter=";"))
        self.assertEqual(len(rows), 2); self.assertEqual(len(rows[0]), 2)


if __name__ == "__main__":
    unittest.main()
